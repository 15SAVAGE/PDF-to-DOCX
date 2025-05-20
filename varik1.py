import os
import pytesseract
import sys
import logging
import tempfile
from PIL import Image

def resource_path(relative_path):
    if hasattr(sys, '_MEIPASS'):
        return os.path.join(sys._MEIPASS, relative_path)
    return os.path.join(os.path.abspath("."), relative_path)

tesseract_path = resource_path("tesseract.exe")
tessdata_path = resource_path("tessdata")

pytesseract.pytesseract.tesseract_cmd = tesseract_path
os.environ['TESSDATA_PREFIX'] = tessdata_path
os.environ['TESSDATA_PREFIX'] = tessdata_path



import io
import fitz
import pytesseract
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
import pikepdf
from pdf2docx import Converter
from PyQt6.QtWidgets import QApplication, QWidget, QPushButton, QVBoxLayout, QLineEdit, QLabel, QFileDialog, QMessageBox, QInputDialog, QProgressBar
import pdfplumber
def is_math_formula(text):
    math_patterns = [
        r'\d+\s*=\s*\d+',
        r'\\[a-zA-Z]+\{',
        r'\\[a-zA-Z]+',
        r'[\^_]\s*\{',
        r'[\^_]\s*[a-zA-Z0-9]',
        r'[α-ωΑ-ΩδΔξΞηΗ]',
        r'\d+\s*/\s*\d+',
        r'[∑∫∏√∞≈≠≡≤≥±→×÷]',
        r'[=+\-*/]',
        r'[ΑΒΓΔΕΖΗΘΙΚΛΜΝΞΠΡΣΤΥΦΧΨΩ\(\)\[\]\{\}⟨⟩⌊⌋⌈⌉⁰¹²³⁴⁵⁶⁷⁸⁹₀₁₂₃₄₅₆₇₈₉−÷≈≠≡≤≥±→√∞∂∆∇∏∑∫]'
    ]
    text = text.strip()
    if len(text) > 300:
         return False
    return any(re.search(pattern, text) for pattern in math_patterns)
def format_math_formula(text):
    replacements = {
        r'\\alpha': 'α',
        r'\\beta': 'β',
        r'\\delta': 'δ',
        r'\\xi': 'ξ',
        r'\\eta': 'η',
        r'\\frac\s*\{([^}]+)\}\s*\{([^}]+)\}': r'\1/\2',
        r'\^\{([^}]+)\}': r'^\1',
        r'_\{([^}]+)\}': r'_\1',
        r'\\cdot': '·',
        r'\\times': '×',
        r'\\approx': '≈',
        r'\\neq': '≠',
        r'\\leq': '≤',
        r'\\geq': '≥',
        r'\\Alpha': 'Α',
        r'\\Beta': 'Β',
        r'\\Gamma': 'Γ',
        r'\\Delta': 'Δ',
        r'\\Epsilon': 'Ε',
        r'\\Zeta': 'Ζ',
        r'\\Eta': 'Η',
        r'\\Theta': 'Θ',
        r'\\Iota': 'Ι',
        r'\\Kappa': 'Κ',
        r'\\Lambda': 'Λ',
        r'\\Mu': 'Μ',
        r'\\Nu': 'Ν',
        r'\\Xi': 'Ξ',
        r'\\Pi': 'Π',
        r'\\Rho': 'Ρ',
        r'\\Sigma': 'Σ',
        r'\\Tau': 'Τ',
        r'\\Upsilon': 'Υ',
        r'\\Phi': 'Φ',
        r'\\Chi': 'Χ',
        r'\\Psi': 'Ψ',
        r'\\Omega': 'Ω',
        r'\\prime': '′',
        r'\\doubleprime': '″',
        r'\\tripleprime': '‴',
        r'\\degree': '°',
        r'\\hbar': 'ℏ',
        r'\\Im': 'ℑ',
        r'\\Re': 'ℜ',
        r'\\wp': '℘',
        r'\\leftarrow': '←',
        r'\\rightarrow': '→',
        r'\\uparrow': '↑',
        r'\\downarrow': '↓',
        r'\\leftrightarrow': '↔',
        r'\\updownarrow': '↕',
        r'\\Leftarrow': '⇐',
        r'\\Rightarrow': '⇒',
        r'\\Leftrightarrow': '⇔',
        r'\\angle': '∠',
        r'\\measuredangle': '∡',
        r'\\parallel': '∥',
        r'\\nparallel': '∦',
        r'\\Perp': '⊥',
        r'\\in': '∈',
        r'\\notin': '∉',
        r'\\subset': '⊂',
        r'\\supset': '⊃',
        r'\\cup': '∪',
        r'\\cap': '∩',
        r'\\emptyset': '∅',
        r'\\aleph': 'ℵ',
        r'\\land': '∧',
        r'\\lor': '∨',
        r'\\lnot': '¬',
        r'\\forall': '∀',
        r'\\exists': '∃',
        r'\\therefore': '∴',
        r'\\because': '∵',
        r'\\frac12': '½',
        r'\\frac13': '⅓',
        r'\\frac14': '¼',
        r'\\frac34': '¾',
        r'\\colon': '∶',
        r'\\vcentcolon': '∷',
        r'\\propto': '∝',
        r'\\sim': '∼',
        r'(?<=\d)-(?=\d)': '−',
        r'(?<=\d)/(?=\d)': '÷'
    }
    for pattern, repl in replacements.items():
        text = re.sub(pattern, repl, text)
    return text.strip()
def add_formatted_paragraph(doc, text):
    if is_math_formula(text):
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run(format_math_formula(text))
        run.font.name = 'Cambria Math'
        run.font.size = Pt(12)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Cambria Math')
    else:
        doc.add_paragraph(text)
def convert_pdf_to_docx(pdf_path, docx_path, pdf_password=""):
    try:
        if pdf_password:
            unlocked_pdf_path = os.path.join(tempfile.gettempdir(), f"unlocked_{os.path.basename(pdf_path)}")
            if not unlock_pdf(pdf_path, unlocked_pdf_path, pdf_password):
                raise Exception("Failed to unlock PDF. Check the password.")
            pdf_path = unlocked_pdf_path

        cv = Converter(pdf_path)
        cv.convert(docx_path)
        cv.close()

    except Exception as e:
        raise Exception(f"Error converting PDF to DOCX: {e}")

    finally:
        if pdf_password and os.path.exists(unlocked_pdf_path):
            os.remove(unlocked_pdf_path)

def clean_ocr_text(text):
    lines = text.split('\n')
    cleaned_lines = []
    for i, line in enumerate(lines):
        if i > 0 and lines[i - 1].endswith('-'):
            cleaned_lines[-1] = cleaned_lines[-1][:-1] + line.strip()
        else:
            cleaned_lines.append(line.strip())
    return '\n'.join(cleaned_lines)

def is_non_standard_font(page):
    fonts = page.get_fonts(full=True)
    for font in fonts:
        font_name = font[3].lower() if len(font) > 3 else ""
        if "bookman" in font_name or "italic" in font_name:  
            return True
    return False

def sanitize_table_row(row):
    return [cell if cell is not None else "" for cell in row]

def extract_table_with_pdfplumber(pdf_path, page_number):
    try:
        with pdfplumber.open(pdf_path) as pdf:
            page = pdf.pages[page_number]
            tables = page.extract_tables()
            if not tables:
                print(f"[INFO] No tables found on page {page_number + 1}.")
                return None
            sanitized_tables = [[sanitize_table_row(row) for row in table] for table in tables]
            return sanitized_tables
    except Exception as e:
        print(f"[ERROR] Error extracting table with pdfplumber: {e}")
        return None

def extract_text_with_ocr(image):
    try:
        return pytesseract.image_to_string(image, lang='rus+eng')
    except pytesseract.TesseractError as e:
        print(f"[ERROR] Tesseract error: {e}")
        return ""
    except Exception as e:
        print(f"[ERROR] OCR error: {e}")
        return ""

def save_image_to_tempfile(img):
    try:
        fd, temp_path = tempfile.mkstemp(suffix=".png")
        os.close(fd) 
        img.save(temp_path)
        return temp_path
    except Exception as e:
        print(f"[ERROR] Error saving image to temp file: {e}")
        return None

def convert_pdf_to_docx_with_ocr(pdf_path, docx_path, pdf_password="", chunk_size=10):
    temp_docx_path = None 
    try:
        if not os.path.exists(pytesseract.pytesseract.tesseract_cmd):
            raise FileNotFoundError(f"Tesseract executable not found at: {pytesseract.pytesseract.tesseract_cmd}")

        required_languages = ['rus.traineddata', 'eng.traineddata']
        for lang in required_languages:
            lang_path = os.path.join(os.environ['TESSDATA_PREFIX'], lang)
            if not os.path.exists(lang_path):
                raise FileNotFoundError(f"Language data file not found: {lang_path}")

        if pdf_password:
            unlocked_pdf_path = os.path.join(tempfile.gettempdir(), f"unlocked_{os.path.basename(pdf_path)}")
            if not unlock_pdf(pdf_path, unlocked_pdf_path, pdf_password):
                raise Exception("Failed to unlock PDF. Check the password.")
            pdf_path = unlocked_pdf_path

        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_docx:
            temp_docx_path = temp_docx.name

        print("[INFO] Converting PDF to DOCX using pdf2docx...")
        cv = Converter(pdf_path)
        cv.convert(temp_docx_path, start=0, end=None)
        cv.close()

        pdf = fitz.open(pdf_path)
        doc = Document(temp_docx_path)

        for page_number in range(len(pdf)):
            logging.info(f"Processing page {page_number + 1}")
            page = pdf[page_number]

            tables = extract_table_with_pdfplumber(pdf_path, page_number)
            if tables:
                logging.info(f"Tables found on page {page_number + 1}")
                doc.add_heading(f"Page {page_number + 1} (Tables)", level=2)
                for table in tables:
                    for row in table:

                        doc.add_paragraph('\t'.join(row))
                continue  

            pix = page.get_pixmap(dpi=300)
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            ocr_text = extract_text_with_ocr(img)
            if ocr_text:
                doc.add_heading(f"Page {page_number + 1} (OCR)", level=2)
                for line in ocr_text.split('\n'):
                    doc.add_paragraph(line.strip())
            else:
                logging.warning(f"OCR could not extract text from page {page_number + 1}")
                print(f"[INFO] Inserting image of page {page_number + 1}...")
                temp_img_path = save_image_to_tempfile(img)
                if temp_img_path:
                    try:
                        doc.add_picture(temp_img_path, width=Inches(6))
                    finally:
                        if os.path.exists(temp_img_path):
                            os.remove(temp_img_path)

        if os.path.exists(docx_path) and not os.access(docx_path, os.W_OK):
            raise PermissionError(f"No write access to file: {docx_path}")
        print(f"[INFO] Saving document to: {docx_path}")
        doc.save(docx_path)

    except PermissionError as e:
        raise Exception(f"Error: {e}")
    except Exception as e:
        raise Exception(f"Error converting PDF to DOCX with OCR: {str(e)}")
    finally:
        if temp_docx_path and os.path.exists(temp_docx_path):
            os.remove(temp_docx_path)
        if pdf_password and os.path.exists(unlocked_pdf_path):
            os.remove(unlocked_pdf_path)


def extract_table_from_ocr(ocr_data):
    rows = []
    current_row = []
    last_top = None

    for i, text in enumerate(ocr_data['text']):
        if text.strip():
            top = ocr_data['top'][i]
            if last_top is None or abs(top - last_top) < 10:
                current_row.append(text)
            else:
                rows.append(current_row)
                current_row = [text]
            last_top = top

    if current_row:
        rows.append(current_row)

    return rows


def unlock_pdf(input_path, output_path, password):
    try:
        with pikepdf.open(input_path, password=password) as pdf:
            pdf.save(output_path)
        print(f"PDF unlocked and saved to: {output_path}")
        return True
    except pikepdf._qpdf.PasswordError:
        print("Incorrect password for PDF.")
        return False
    except Exception as e:
        print(f"Error unlocking PDF: {e}")
        return False


class ConverterGUI(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("PDF to DOCX Converter")
        self.setGeometry(300, 300, 400, 300)
        self.pdf_file = ""
        self.output_folder = ""
        layout = QVBoxLayout()
        self.pdf_label = QLabel("Selected PDF: not selected")
        self.folder_label = QLabel("Selected folder: not selected")
        layout.addWidget(self.pdf_label)
        layout.addWidget(self.folder_label)
        self.btn_select_pdf = QPushButton("Select PDF file")
        self.btn_select_pdf.clicked.connect(self.select_pdf)
        layout.addWidget(self.btn_select_pdf)
        self.btn_select_folder = QPushButton("Select output folder")
        self.btn_select_folder.clicked.connect(self.select_folder)
        layout.addWidget(self.btn_select_folder)
        self.btn_convert = QPushButton("Convert")
        self.btn_convert.clicked.connect(self.convert)
        layout.addWidget(self.btn_convert)

        self.progress_bar = QProgressBar()
        self.progress_bar.setValue(0)
        layout.addWidget(self.progress_bar)
        self.console_output = QLabel("Waiting to start...")
        layout.addWidget(self.console_output)

        self.setLayout(layout)

    def update_progress(self, value, console_text=None):
        self.progress_bar.setValue(value)
        if console_text:
            self.console_output.setText(console_text)
        QApplication.processEvents() 

    def reset_progress(self):
        self.progress_bar.setValue(0)
        self.console_output.setText("Waiting to start...")

    def select_pdf(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "Select PDF file", "", "PDF Files (*.pdf)")
        if file_path:
            self.pdf_file = file_path
            self.pdf_label.setText(f"Selected PDF: {self.pdf_file}")

    def select_folder(self):
        folder = QFileDialog.getExistingDirectory(self, "Select output folder")
        if folder:
            self.output_folder = folder
            self.folder_label.setText(f"Selected folder: {self.output_folder}")

    def convert(self):
        if not self.pdf_file:
            QMessageBox.warning(self, "Warning", "Please select a PDF file first.")
            return
        if not self.output_folder:
            QMessageBox.warning(self, "Warning", "Please select a folder for output.")
            return
        try:
            self.reset_progress()
            self.update_progress(0, "Starting processing...")
            basename = os.path.splitext(os.path.basename(self.pdf_file))[0]
            docx_path = os.path.join(self.output_folder, f"{basename}.docx")
            docx_path = os.path.normpath(docx_path)

            if not os.path.isdir(self.output_folder):
                raise Exception(f"Folder does not exist or is not accessible: {self.output_folder}")
            if not os.access(self.output_folder, os.W_OK):
                raise Exception(f"No write access to folder: {self.output_folder}")

            unlocked_pdf_path = None
            try:
                pdf = fitz.open(self.pdf_file)
                if pdf.is_encrypted:
                    password, ok = QInputDialog.getText(self, "Protected PDF", "Enter password for PDF:", QLineEdit.Password)
                    if not ok or not password:
                        QMessageBox.warning(self, "Warning", "Password not entered. Conversion cancelled.")
                        return
                    unlocked_pdf_path = os.path.join(tempfile.gettempdir(), f"unlocked_{basename}.pdf")
                    if not unlock_pdf(self.pdf_file, unlocked_pdf_path, password):
                        QMessageBox.critical(self, "Error", "Failed to unlock PDF. Check the password.")
                        return
                    self.pdf_file = unlocked_pdf_path
            except Exception as e:
                self.update_progress(0, f"[ERROR] {e}")
                return

            if os.path.exists(docx_path):
                docx_path = os.path.join(self.output_folder, f"New_{os.path.basename(docx_path)}")

            pdf = fitz.open(self.pdf_file)
            total_pages = len(pdf)
            for i in range(total_pages):
                progress = int((i + 1) / total_pages * 100)
                self.update_progress(progress, f"Processing page {i + 1} of {total_pages}")
                QApplication.processEvents()  

                

            self.update_progress(100, "Saving document and changing format...")
            convert_pdf_to_docx_with_ocr(self.pdf_file, docx_path)
            self.update_progress(100, "Conversion completed.")
            QMessageBox.information(self, "Success", f"Conversion completed.\nFile saved: {docx_path}")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error:\n{str(e)}")
        finally:
            self.reset_progress()
            if unlocked_pdf_path and os.path.exists(unlocked_pdf_path):
                os.remove(unlocked_pdf_path)

if __name__ == "__main__":
    app = QApplication(sys.argv)
    gui = ConverterGUI()
    gui.show()
    sys.exit(app.exec())